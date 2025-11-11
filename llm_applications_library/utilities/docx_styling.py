"""
DOCX document styling utilities for professional document formatting
"""

import logging
from typing import ClassVar

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

logger = logging.getLogger(__name__)


class DocxStyler:
    """Professional DOCX document styling utility"""

    # カラーテーマ定義
    COLOR_THEMES: ClassVar = {
        "blue": {
            "primary": RGBColor(68, 114, 196),  # #4472C4
            "secondary": RGBColor(142, 169, 219),  # #8EA9DB
            "accent": RGBColor(217, 225, 242),  # #D9E1F2
            "text": RGBColor(44, 62, 80),  # #2C3E50
        },
        "green": {
            "primary": RGBColor(112, 173, 71),  # #70AD47
            "secondary": RGBColor(169, 208, 142),  # #A9D08E
            "accent": RGBColor(226, 239, 218),  # #E2EFDA
            "text": RGBColor(39, 78, 19),  # #274E13
        },
        "orange": {
            "primary": RGBColor(255, 192, 0),  # #FFC000
            "secondary": RGBColor(255, 217, 102),  # #FFD966
            "accent": RGBColor(255, 242, 204),  # #FFF2CC
            "text": RGBColor(127, 96, 0),  # #7F6000
        },
        "purple": {
            "primary": RGBColor(112, 48, 160),  # #7030A0
            "secondary": RGBColor(174, 125, 209),  # #AE7DD1
            "accent": RGBColor(230, 218, 241),  # #E6DAF1
            "text": RGBColor(56, 24, 80),  # #381850
        },
        "gray": {
            "primary": RGBColor(89, 89, 89),  # #595959
            "secondary": RGBColor(166, 166, 166),  # #A6A6A6
            "accent": RGBColor(217, 217, 217),  # #D9D9D9
            "text": RGBColor(64, 64, 64),  # #404040
        },
    }

    # プロフェッショナルテーブルスタイル定義
    TABLE_STYLES: ClassVar = {
        "professional": {
            "style_id": "LightShading-Accent1",
            "name": "Light Shading - Accent 1",
        },
        "elegant": {
            "style_id": "MediumShading1-Accent1",
            "name": "Medium Shading 1 - Accent 1",
        },
        "modern": {"style_id": "LightList-Accent1", "name": "Light List - Accent 1"},
        "minimal": {"style_id": "TableGrid", "name": "Table Grid"},
        "colorful": {
            "style_id": "ColorfulShading-Accent1",
            "name": "Colorful Shading - Accent 1",
        },
    }

    def __init__(self, document: Document):
        self.document = document

    def apply_document_theme(
        self,
        theme: str = "professional",
        font_name: str = "Calibri",
        font_size: int = 11,
        color_theme: str = "blue",
        line_spacing: float = 1.15,
        paragraph_spacing: int = 6,
    ):
        """ドキュメント全体のテーマを適用"""
        logger.info(f"Applying document theme: {theme} with font: {font_name}")

        try:
            # 基本スタイルの設定
            self._set_default_font(font_name, font_size)
            self._set_paragraph_spacing(line_spacing, paragraph_spacing)

            # 見出しスタイルの設定
            self._configure_heading_styles(color_theme, font_name)

            # ページ設定の調整
            self._configure_page_settings(theme)

        except Exception as e:
            logger.error(f"Error applying document theme: {e}")

    def _set_default_font(self, font_name: str, font_size: int):
        """デフォルトフォントの設定"""
        try:
            normal_style = self.document.styles["Normal"]
            normal_font = normal_style.font
            normal_font.name = font_name
            normal_font.size = Pt(font_size)

            # 段落設定
            paragraph_format = normal_style.paragraph_format
            paragraph_format.space_after = Pt(6)
            paragraph_format.line_spacing = 1.15

        except Exception as e:
            logger.warning(f"Could not set default font: {e}")

    def _set_paragraph_spacing(self, line_spacing: float, paragraph_spacing: int):
        """段落間隔の設定"""
        try:
            normal_style = self.document.styles["Normal"]
            paragraph_format = normal_style.paragraph_format
            paragraph_format.line_spacing = line_spacing
            paragraph_format.space_after = Pt(paragraph_spacing)
        except Exception as e:
            logger.warning(f"Could not set paragraph spacing: {e}")

    def _configure_heading_styles(self, color_theme: str, font_name: str):
        """見出しスタイルの設定"""
        try:
            colors = self.COLOR_THEMES.get(color_theme, self.COLOR_THEMES["blue"])

            # 見出し1-3のスタイル設定
            heading_configs = [
                ("Heading 1", 16, True, colors["primary"]),
                ("Heading 2", 14, True, colors["secondary"]),
                ("Heading 3", 12, True, colors["text"]),
            ]

            for style_name, size, bold, color in heading_configs:
                try:
                    style = self.document.styles[style_name]
                    font = style.font
                    font.name = font_name
                    font.size = Pt(size)
                    font.bold = bold
                    font.color.rgb = color

                    # 見出しの段落設定
                    paragraph_format = style.paragraph_format
                    paragraph_format.space_before = Pt(12)
                    paragraph_format.space_after = Pt(6)

                except KeyError:
                    logger.warning(f"Heading style {style_name} not found")

        except Exception as e:
            logger.error(f"Error configuring heading styles: {e}")

    def _configure_page_settings(self, theme: str):
        """ページ設定の調整"""
        try:
            sections = self.document.sections
            for section in sections:
                # Configure margins based on theme
                if theme == "professional":
                    section.top_margin = Inches(1.0)
                    section.bottom_margin = Inches(1.0)
                    section.left_margin = Inches(1.0)
                    section.right_margin = Inches(1.0)
                elif theme == "modern":
                    section.top_margin = Inches(0.8)
                    section.bottom_margin = Inches(0.8)
                    section.left_margin = Inches(0.8)
                    section.right_margin = Inches(0.8)
                elif theme == "minimal":
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)

        except Exception as e:
            logger.warning(f"Could not configure page settings: {e}")

    def style_tables(
        self,
        table_style_name: str = "professional",
        header_bg_color: str | None = None,
        header_text_color: str | None = None,
        alternating_rows: bool = True,
        _border_style: str = "medium",
    ):
        """ドキュメント内の全テーブルにスタイルを適用"""
        logger.info(f"Styling tables with style: {table_style_name}")

        try:
            tables = self.document.tables
            for table in tables:
                self._apply_table_style(
                    table,
                    table_style_name,
                    header_bg_color,
                    header_text_color,
                    alternating_rows,
                    _border_style,
                )

        except Exception as e:
            logger.error(f"Error styling tables: {e}")

    def _apply_table_style(
        self,
        table: Table,
        style_name: str,
        header_bg_color: str | None = None,
        header_text_color: str | None = None,
        alternating_rows: bool = True,
        _border_style: str = "medium",
    ):
        """個別テーブルへのスタイル適用"""
        try:
            # Apply Word built-in styles
            try:
                if style_name == "professional":
                    table.style = "Light Shading - Accent 1"
                elif style_name == "elegant":
                    table.style = "Medium Shading 1 - Accent 1"
                elif style_name == "modern":
                    table.style = "Light List - Accent 1"
                elif style_name == "minimal":
                    table.style = "Table Grid"
                elif style_name == "colorful":
                    # このスタイルが利用できない場合はTable Gridを使用
                    try:
                        table.style = "Colorful Shading - Accent 1"
                    except KeyError:
                        logger.warning(
                            "Colorful Shading style not available, using Table Grid"
                        )
                        table.style = "Table Grid"
                else:
                    table.style = "Light Shading - Accent 1"
            except Exception as style_error:
                logger.warning(f"Could not apply style {style_name}: {style_error}")
                # フォールバックとしてTable Gridを使用
                table.style = "Table Grid"

            # テーブル配置
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # カスタムヘッダースタイル適用
            if header_bg_color or header_text_color:
                self._apply_header_styling(table, header_bg_color, header_text_color)

            # 行の交互背景色
            if alternating_rows and style_name not in ["colorful"]:
                self._apply_alternating_rows(table)

        except Exception as e:
            logger.warning(f"Could not apply table style {style_name}: {e}")

    def _apply_header_styling(
        self, table: Table, bg_color: str | None, text_color: str | None
    ):
        """ヘッダー行のカスタムスタイリング"""
        try:
            if table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    # 背景色設定
                    if bg_color:
                        self._set_cell_background_color(cell, bg_color)

                    # 文字色とフォント設定
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if text_color:
                                run.font.color.rgb = RGBColor.from_string(
                                    text_color.replace("#", "")
                                )
                            run.font.bold = True

        except Exception as e:
            logger.warning(f"Could not apply header styling: {e}")

    def _apply_alternating_rows(self, table: Table):
        """行の交互背景色を適用"""
        try:
            for i, row in enumerate(table.rows[1:], 1):  # ヘッダー行をスキップ
                if i % 2 == 0:  # 偶数行
                    for cell in row.cells:
                        self._set_cell_background_color(cell, "#F8F9FA")

        except Exception as e:
            logger.warning(f"Could not apply alternating rows: {e}")

    def _set_cell_background_color(self, cell, color: str):
        """セルの背景色を設定"""
        try:
            # Set background color via XML
            cell_xml = cell._element  # noqa: SLF001
            table_cell_properties = cell_xml.get_or_add_tcPr()

            # 既存のshading要素を削除
            for shading in table_cell_properties.xpath(".//w:shd"):
                shading.getparent().remove(shading)

            # 新しいshading要素を作成
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), color.replace("#", ""))
            table_cell_properties.append(shading)

        except Exception as e:
            logger.warning(f"Could not set cell background color: {e}")

    def add_header_footer(
        self,
        add_header: bool = False,
        add_footer: bool = False,
        header_text: str = "",
        footer_text: str = "",
    ):
        """ヘッダー・フッター追加"""
        try:
            if add_header and header_text:
                header = self.document.sections[0].header
                header_para = header.paragraphs[0]
                header_para.text = header_text
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if add_footer and footer_text:
                footer = self.document.sections[0].footer
                footer_para = footer.paragraphs[0]
                footer_para.text = footer_text
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            logger.warning(f"Could not add header/footer: {e}")


def apply_professional_styling(
    document: Document,
    document_style: dict | None = None,
    table_style: dict | None = None,
    add_header: bool = False,
    add_footer: bool = False,
) -> Document:
    """ドキュメントにプロフェッショナルなスタイリングを適用"""

    styler = DocxStyler(document)

    # デフォルトスタイル設定
    doc_style = document_style or {}
    theme = doc_style.get("document_theme", "professional")
    font_name = doc_style.get("font_name", "Calibri")
    font_size = doc_style.get("font_size", 11)
    color_theme = doc_style.get("color_theme", "blue")
    line_spacing = doc_style.get("line_spacing", 1.15)
    paragraph_spacing = doc_style.get("paragraph_spacing", 6)

    # ドキュメントテーマ適用
    styler.apply_document_theme(
        theme, font_name, font_size, color_theme, line_spacing, paragraph_spacing
    )

    # テーブルスタイル適用
    if document.tables:
        tbl_style = table_style or {}
        styler.style_tables(
            table_style_name=tbl_style.get("style_name", "professional"),
            header_bg_color=tbl_style.get("header_bg_color"),
            header_text_color=tbl_style.get("header_text_color"),
            alternating_rows=tbl_style.get("alternating_rows", True),
        )

    # ヘッダー・フッター追加
    if add_header or add_footer:
        styler.add_header_footer(
            add_header=add_header,
            add_footer=add_footer,
            header_text="Professional Document",
            footer_text="Generated with Professional DOCX Converter",
        )

    return document
