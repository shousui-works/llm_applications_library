"""Tests for docx_styling utilities."""

import pytest
from unittest.mock import Mock, patch

# Mock docx module if not available
try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = Mock
    WD_TABLE_ALIGNMENT = Mock()
    WD_TABLE_ALIGNMENT.CENTER = "center"
    WD_ALIGN_PARAGRAPH = Mock()
    WD_ALIGN_PARAGRAPH.CENTER = "center"
    Inches = Mock(side_effect=lambda x: f"inches_{x}")
    Pt = Mock(side_effect=lambda x: f"pt_{x}")

    class MockRGBColor:
        def __init__(self, r, g, b):
            self.r, self.g, self.b = r, g, b

        @classmethod
        def from_string(cls, hex_str):
            return cls(0, 0, 0)

    RGBColor = MockRGBColor

from llm_applications_library.utilities.docx_styling import (
    DocxStyler,
    apply_professional_styling,
)


class TestDocxStyler:
    """Test cases for DocxStyler class."""

    @pytest.fixture
    def mock_document(self):
        """Create a mock Document object."""
        document = Mock(spec=Document)
        document.styles = {}
        document.tables = []
        document.sections = []
        return document

    @pytest.fixture
    def docx_styler(self, mock_document):
        """Create DocxStyler instance with mock document."""
        return DocxStyler(mock_document)

    def test_init(self, mock_document):
        """Test DocxStyler initialization."""
        styler = DocxStyler(mock_document)
        assert styler.document == mock_document

    def test_color_themes_exist(self):
        """Test that predefined color themes exist."""
        assert "blue" in DocxStyler.COLOR_THEMES
        assert "green" in DocxStyler.COLOR_THEMES
        assert "orange" in DocxStyler.COLOR_THEMES
        assert "purple" in DocxStyler.COLOR_THEMES
        assert "gray" in DocxStyler.COLOR_THEMES

        # Test blue theme structure
        blue_theme = DocxStyler.COLOR_THEMES["blue"]
        assert "primary" in blue_theme
        assert "secondary" in blue_theme
        assert "accent" in blue_theme
        assert "text" in blue_theme

        # Test color types
        assert isinstance(blue_theme["primary"], RGBColor)

    def test_table_styles_exist(self):
        """Test that predefined table styles exist."""
        assert "professional" in DocxStyler.TABLE_STYLES
        assert "elegant" in DocxStyler.TABLE_STYLES
        assert "modern" in DocxStyler.TABLE_STYLES
        assert "minimal" in DocxStyler.TABLE_STYLES
        assert "colorful" in DocxStyler.TABLE_STYLES

    def test_apply_document_theme_default(self, docx_styler):
        """Test applying default document theme."""
        with (
            patch.object(docx_styler, "_set_default_font") as mock_font,
            patch.object(docx_styler, "_set_paragraph_spacing") as mock_spacing,
            patch.object(docx_styler, "_configure_heading_styles") as mock_headings,
            patch.object(docx_styler, "_configure_page_settings") as mock_page,
        ):
            docx_styler.apply_document_theme()

            mock_font.assert_called_once_with("Calibri", 11)
            mock_spacing.assert_called_once_with(1.15, 6)
            mock_headings.assert_called_once_with("blue", "Calibri")
            mock_page.assert_called_once_with("professional")

    def test_apply_document_theme_custom(self, docx_styler):
        """Test applying custom document theme."""
        with (
            patch.object(docx_styler, "_set_default_font") as mock_font,
            patch.object(docx_styler, "_set_paragraph_spacing") as mock_spacing,
            patch.object(docx_styler, "_configure_heading_styles") as mock_headings,
            patch.object(docx_styler, "_configure_page_settings") as mock_page,
        ):
            docx_styler.apply_document_theme(
                theme="modern",
                font_name="Arial",
                font_size=12,
                color_theme="green",
                line_spacing=1.2,
                paragraph_spacing=8,
            )

            mock_font.assert_called_once_with("Arial", 12)
            mock_spacing.assert_called_once_with(1.2, 8)
            mock_headings.assert_called_once_with("green", "Arial")
            mock_page.assert_called_once_with("modern")

    def test_set_default_font_success(self, docx_styler):
        """Test successful font setting."""
        mock_style = Mock()
        mock_font = Mock()
        mock_paragraph_format = Mock()

        mock_style.font = mock_font
        mock_style.paragraph_format = mock_paragraph_format
        docx_styler.document.styles = {"Normal": mock_style}

        docx_styler._set_default_font("Arial", 12)

        assert mock_font.name == "Arial"
        assert mock_font.size == Pt(12)
        assert mock_paragraph_format.space_after == Pt(6)
        assert mock_paragraph_format.line_spacing == 1.15

    def test_set_default_font_error_handling(self, docx_styler):
        """Test font setting error handling."""
        docx_styler.document.styles = {}

        # Should not raise exception
        docx_styler._set_default_font("Arial", 12)

    def test_set_paragraph_spacing_success(self, docx_styler):
        """Test successful paragraph spacing setting."""
        mock_style = Mock()
        mock_paragraph_format = Mock()
        mock_style.paragraph_format = mock_paragraph_format
        docx_styler.document.styles = {"Normal": mock_style}

        docx_styler._set_paragraph_spacing(1.5, 10)

        assert mock_paragraph_format.line_spacing == 1.5
        assert mock_paragraph_format.space_after == Pt(10)

    def test_configure_heading_styles_success(self, docx_styler):
        """Test successful heading styles configuration."""
        mock_styles = {}
        for heading in ["Heading 1", "Heading 2", "Heading 3"]:
            mock_style = Mock()
            mock_font = Mock()
            mock_color = Mock()
            mock_paragraph_format = Mock()

            mock_style.font = mock_font
            mock_font.color = mock_color
            mock_style.paragraph_format = mock_paragraph_format
            mock_styles[heading] = mock_style

        docx_styler.document.styles = mock_styles

        docx_styler._configure_heading_styles("blue", "Arial")

        # Check that all heading styles were configured
        for heading in ["Heading 1", "Heading 2", "Heading 3"]:
            style = mock_styles[heading]
            assert style.font.name == "Arial"
            assert style.font.bold is True

    def test_configure_page_settings_professional(self, docx_styler):
        """Test professional page settings configuration."""
        mock_section = Mock()
        docx_styler.document.sections = [mock_section]

        docx_styler._configure_page_settings("professional")

        assert mock_section.top_margin == Inches(1.0)
        assert mock_section.bottom_margin == Inches(1.0)
        assert mock_section.left_margin == Inches(1.0)
        assert mock_section.right_margin == Inches(1.0)

    def test_configure_page_settings_modern(self, docx_styler):
        """Test modern page settings configuration."""
        mock_section = Mock()
        docx_styler.document.sections = [mock_section]

        docx_styler._configure_page_settings("modern")

        assert mock_section.top_margin == Inches(0.8)
        assert mock_section.bottom_margin == Inches(0.8)
        assert mock_section.left_margin == Inches(0.8)
        assert mock_section.right_margin == Inches(0.8)

    def test_configure_page_settings_minimal(self, docx_styler):
        """Test minimal page settings configuration."""
        mock_section = Mock()
        docx_styler.document.sections = [mock_section]

        docx_styler._configure_page_settings("minimal")

        assert mock_section.top_margin == Inches(0.5)
        assert mock_section.bottom_margin == Inches(0.5)
        assert mock_section.left_margin == Inches(0.5)
        assert mock_section.right_margin == Inches(0.5)

    def test_style_tables(self, docx_styler):
        """Test table styling."""
        mock_table = Mock()
        docx_styler.document.tables = [mock_table]

        with patch.object(docx_styler, "_apply_table_style") as mock_apply:
            docx_styler.style_tables("professional")
            mock_apply.assert_called_once_with(
                mock_table, "professional", None, None, True, "medium"
            )

    def test_apply_table_style_professional(self, docx_styler):
        """Test applying professional table style."""
        mock_table = Mock()
        mock_table.style = None

        docx_styler._apply_table_style(mock_table, "professional")

        assert mock_table.style == "Light Shading - Accent 1"
        assert mock_table.alignment == WD_TABLE_ALIGNMENT.CENTER

    def test_apply_table_style_elegant(self, docx_styler):
        """Test applying elegant table style."""
        mock_table = Mock()
        mock_table.style = None

        docx_styler._apply_table_style(mock_table, "elegant")

        assert mock_table.style == "Medium Shading 1 - Accent 1"

    def test_apply_table_style_modern(self, docx_styler):
        """Test applying modern table style."""
        mock_table = Mock()
        mock_table.style = None

        docx_styler._apply_table_style(mock_table, "modern")

        assert mock_table.style == "Light List - Accent 1"

    def test_apply_table_style_minimal(self, docx_styler):
        """Test applying minimal table style."""
        mock_table = Mock()
        mock_table.style = None

        docx_styler._apply_table_style(mock_table, "minimal")

        assert mock_table.style == "Table Grid"

    def test_apply_table_style_colorful_fallback(self, docx_styler):
        """Test colorful table style with fallback."""
        mock_table = Mock()

        def style_setter(self, value):
            if value == "Colorful Shading - Accent 1":
                raise KeyError("Style not available")
            self._style = value

        type(mock_table).style = property(
            lambda x: getattr(x, "_style", None), style_setter
        )

        docx_styler._apply_table_style(mock_table, "colorful")

        assert mock_table._style == "Table Grid"

    def test_apply_header_styling(self, docx_styler):
        """Test header styling application."""
        mock_cell = Mock()
        mock_paragraph = Mock()
        mock_run = Mock()
        mock_font = Mock()
        mock_color = Mock()

        mock_run.font = mock_font
        mock_font.color = mock_color
        mock_paragraph.runs = [mock_run]
        mock_cell.paragraphs = [mock_paragraph]

        mock_row = Mock()
        mock_row.cells = [mock_cell]
        mock_table = Mock()
        mock_table.rows = [mock_row]

        with patch.object(docx_styler, "_set_cell_background_color") as mock_bg:
            docx_styler._apply_header_styling(mock_table, "#FFFFFF", "#000000")

            mock_bg.assert_called_once_with(mock_cell, "#FFFFFF")
            assert mock_font.bold is True

    def test_apply_alternating_rows(self, docx_styler):
        """Test alternating row colors."""
        mock_cells = [Mock(), Mock()]
        mock_row1 = Mock()
        mock_row1.cells = mock_cells
        mock_row2 = Mock()
        mock_row2.cells = mock_cells
        mock_row3 = Mock()
        mock_row3.cells = mock_cells

        mock_table = Mock()
        mock_table.rows = [Mock(), mock_row1, mock_row2, mock_row3]  # First is header

        with patch.object(docx_styler, "_set_cell_background_color") as mock_bg:
            docx_styler._apply_alternating_rows(mock_table)

            # Should color even rows (mock_row2)
            assert mock_bg.call_count == 2  # 2 cells in even row

    def test_set_cell_background_color(self, docx_styler):
        """Test cell background color setting."""
        mock_element = Mock()
        mock_tc_pr = Mock()
        mock_element.get_or_add_tcPr.return_value = mock_tc_pr
        mock_tc_pr.xpath.return_value = []

        mock_cell = Mock()
        mock_cell._element = mock_element

        with (
            patch(
                "llm_applications_library.utilities.docx_styling.OxmlElement"
            ) as mock_oxmlelement,
            patch("llm_applications_library.utilities.docx_styling.qn") as mock_qn,
        ):
            mock_shading = Mock()
            mock_oxmlelement.return_value = mock_shading

            docx_styler._set_cell_background_color(mock_cell, "#FFFFFF")

            mock_oxmlelement.assert_called_once_with("w:shd")
            mock_shading.set.assert_any_call(mock_qn("w:val"), "clear")
            mock_shading.set.assert_any_call(mock_qn("w:color"), "auto")
            mock_shading.set.assert_any_call(mock_qn("w:fill"), "FFFFFF")

    def test_add_header_footer(self, docx_styler):
        """Test adding header and footer."""
        mock_header = Mock()
        mock_footer = Mock()
        mock_header_para = Mock()
        mock_footer_para = Mock()

        mock_header.paragraphs = [mock_header_para]
        mock_footer.paragraphs = [mock_footer_para]

        mock_section = Mock()
        mock_section.header = mock_header
        mock_section.footer = mock_footer
        docx_styler.document.sections = [mock_section]

        docx_styler.add_header_footer(
            add_header=True,
            add_footer=True,
            header_text="Test Header",
            footer_text="Test Footer",
        )

        assert mock_header_para.text == "Test Header"
        assert mock_footer_para.text == "Test Footer"
        assert mock_header_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert mock_footer_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


class TestApplyProfessionalStyling:
    """Test cases for apply_professional_styling function."""

    @pytest.fixture
    def mock_document(self):
        """Create a mock Document object."""
        document = Mock(spec=Document)
        document.tables = []
        return document

    def test_apply_professional_styling_default(self, mock_document):
        """Test apply_professional_styling with default settings."""
        with patch(
            "llm_applications_library.utilities.docx_styling.DocxStyler"
        ) as mock_styler_class:
            mock_styler = Mock()
            mock_styler_class.return_value = mock_styler

            result = apply_professional_styling(mock_document)

            assert result == mock_document
            mock_styler_class.assert_called_once_with(mock_document)
            mock_styler.apply_document_theme.assert_called_once_with(
                "professional", "Calibri", 11, "blue", 1.15, 6
            )

    def test_apply_professional_styling_custom_document_style(self, mock_document):
        """Test apply_professional_styling with custom document style."""
        document_style = {
            "document_theme": "modern",
            "font_name": "Arial",
            "font_size": 12,
            "color_theme": "green",
            "line_spacing": 1.2,
            "paragraph_spacing": 8,
        }

        with patch(
            "llm_applications_library.utilities.docx_styling.DocxStyler"
        ) as mock_styler_class:
            mock_styler = Mock()
            mock_styler_class.return_value = mock_styler

            apply_professional_styling(mock_document, document_style=document_style)

            mock_styler.apply_document_theme.assert_called_once_with(
                "modern", "Arial", 12, "green", 1.2, 8
            )

    def test_apply_professional_styling_with_tables(self, mock_document):
        """Test apply_professional_styling with tables."""
        mock_document.tables = [Mock()]  # Add a table

        table_style = {
            "style_name": "elegant",
            "header_bg_color": "#FFFFFF",
            "header_text_color": "#000000",
            "alternating_rows": False,
        }

        with patch(
            "llm_applications_library.utilities.docx_styling.DocxStyler"
        ) as mock_styler_class:
            mock_styler = Mock()
            mock_styler_class.return_value = mock_styler

            apply_professional_styling(mock_document, table_style=table_style)

            mock_styler.style_tables.assert_called_once_with(
                table_style_name="elegant",
                header_bg_color="#FFFFFF",
                header_text_color="#000000",
                alternating_rows=False,
            )

    def test_apply_professional_styling_with_header_footer(self, mock_document):
        """Test apply_professional_styling with header and footer."""
        with patch(
            "llm_applications_library.utilities.docx_styling.DocxStyler"
        ) as mock_styler_class:
            mock_styler = Mock()
            mock_styler_class.return_value = mock_styler

            apply_professional_styling(mock_document, add_header=True, add_footer=True)

            mock_styler.add_header_footer.assert_called_once_with(
                add_header=True,
                add_footer=True,
                header_text="Professional Document",
                footer_text="Generated with Professional DOCX Converter",
            )

    def test_apply_professional_styling_no_tables(self, mock_document):
        """Test apply_professional_styling with no tables."""
        mock_document.tables = []

        with patch(
            "llm_applications_library.utilities.docx_styling.DocxStyler"
        ) as mock_styler_class:
            mock_styler = Mock()
            mock_styler_class.return_value = mock_styler

            apply_professional_styling(mock_document)

            mock_styler.style_tables.assert_not_called()
